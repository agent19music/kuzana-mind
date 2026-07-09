import csv
import io
import mimetypes
from pathlib import Path


SUPPORTED = {".txt", ".md", ".pdf", ".docx", ".html", ".htm", ".csv"}
MAX_FILE_BYTES = 25 * 1024 * 1024  # 25 MB


def extract_text(filename: str, content: bytes) -> str:
    """
    Extract plain text from file bytes. Returns empty string if unsupported.
    Raises ValueError for files that exceed the size limit.
    """
    if len(content) > MAX_FILE_BYTES:
        raise ValueError(f"{filename} exceeds 25 MB limit ({len(content) // 1_048_576} MB)")

    ext = Path(filename).suffix.lower()

    if ext in (".txt", ".md"):
        return content.decode("utf-8", errors="replace")

    if ext == ".pdf":
        import pdfplumber
        pages = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text.strip())
        return "\n\n".join(pages)

    if ext == ".docx":
        from docx import Document
        doc = Document(io.BytesIO(content))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                # Preserve heading levels as markdown so the splitter can use them
                style = para.style.name.lower()
                if "heading 1" in style:
                    parts.append(f"# {para.text}")
                elif "heading 2" in style:
                    parts.append(f"## {para.text}")
                elif "heading 3" in style:
                    parts.append(f"### {para.text}")
                else:
                    parts.append(para.text)
        # Include tables as simple text
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n\n".join(parts)

    if ext in (".html", ".htm"):
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(content.decode("utf-8", errors="replace"), "html.parser")
        # Remove script and style blocks
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        return soup.get_text(separator="\n", strip=True)

    if ext == ".csv":
        text = content.decode("utf-8", errors="replace")
        reader = csv.DictReader(io.StringIO(text))
        lines = []
        for row in reader:
            line = ", ".join(f"{k}: {v}" for k, v in row.items() if v.strip())
            if line:
                lines.append(line)
        return "\n".join(lines)

    return ""  # unsupported — caller adds to skipped list
